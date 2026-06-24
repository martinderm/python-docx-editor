#!/usr/bin/env python3
"""Convert Markdown to DOCX with compact, practical default formatting.

Usage:
  py -3 scripts/markdown_to_docx.py --in input.md --out output.docx
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


def existing_markdown(path_str: str) -> Path:
    path = Path(path_str)
    if not path.exists():
        raise argparse.ArgumentTypeError(f"Datei nicht gefunden: {path_str}")
    if path.suffix.lower() not in {".md", ".markdown", ".txt"}:
        raise argparse.ArgumentTypeError(f"Keine Markdown-Datei: {path_str}")
    return path


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(18)
    title.font.bold = True
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.line_spacing = 1.0

    for style_name, size, before, after in (
        ("Heading 1", 14, 10, 4),
        ("Heading 2", 12, 8, 3),
        ("Heading 3", 11, 6, 2),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(10.5)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(2)
            style.paragraph_format.line_spacing = 1.0

    if "Markdown Quote" not in doc.styles:
        style = doc.styles.add_style("Markdown Quote", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.italic = True
        style.paragraph_format.left_indent = Cm(0.8)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)


def set_paragraph_numbering(paragraph: Paragraph, num_id: int, ilvl: int = 0) -> None:
    ppr = paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
    if ppr.numPr is not None:
        ppr.remove(ppr.numPr)

    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    ppr.append(num_pr)


def get_abstract_num_id_for_num_id(anchor: Paragraph, num_id: int) -> int | None:
    try:
        root = anchor.part.numbering_part._element  # pylint: disable=protected-access
        for num in root.findall(qn("w:num")):
            num_id_raw = num.get(qn("w:numId"))
            if not num_id_raw or int(num_id_raw) != int(num_id):
                continue
            abs_ref = num.find(qn("w:abstractNumId"))
            if abs_ref is None:
                return None
            abs_id_raw = abs_ref.get(qn("w:val"))
            return int(abs_id_raw) if abs_id_raw is not None else None
    except Exception:
        return None
    return None


def create_numbering_instance(anchor: Paragraph, abstract_num_id: int, *, restart_at_one: bool = False) -> int | None:
    try:
        root = anchor.part.numbering_part._element  # pylint: disable=protected-access
        max_num_id = 0
        for num in root.findall(qn("w:num")):
            raw = num.get(qn("w:numId"))
            if not raw:
                continue
            max_num_id = max(max_num_id, int(raw))

        new_num_id = max_num_id + 1
        num_el = OxmlElement("w:num")
        num_el.set(qn("w:numId"), str(new_num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(int(abstract_num_id)))
        num_el.append(abs_ref)

        if restart_at_one:
            lvl_override = OxmlElement("w:lvlOverride")
            lvl_override.set(qn("w:ilvl"), "0")
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), "1")
            lvl_override.append(start_override)
            num_el.append(lvl_override)

        root.append(num_el)
        return new_num_id
    except Exception:
        return None


def find_numbering_for_kind(anchor: Paragraph, kind: str) -> tuple[int, int] | None:
    wanted = "bullet" if kind == "ul" else "decimal"
    try:
        root = anchor.part.numbering_part._element  # pylint: disable=protected-access
        abstract_meta: dict[str, tuple[str | None, str | None]] = {}
        for abstract in root.findall(qn("w:abstractNum")):
            abs_id = abstract.get(qn("w:abstractNumId"))
            if not abs_id:
                continue
            lvl0 = None
            for lvl in abstract.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl"))
                if ilvl in {"0", None}:
                    lvl0 = lvl
                    break
            if lvl0 is None:
                continue
            num_fmt_el = lvl0.find(qn("w:numFmt"))
            num_fmt_val = num_fmt_el.get(qn("w:val")) if num_fmt_el is not None else None
            lvl_text_el = lvl0.find(qn("w:lvlText"))
            lvl_text_val = lvl_text_el.get(qn("w:val")) if lvl_text_el is not None else None
            abstract_meta[abs_id] = (num_fmt_val, lvl_text_val)

        candidates: list[tuple[int, str, str | None]] = []
        for num in root.findall(qn("w:num")):
            num_id_raw = num.get(qn("w:numId"))
            abs_ref = num.find(qn("w:abstractNumId"))
            if not num_id_raw or abs_ref is None:
                continue
            abs_id = abs_ref.get(qn("w:val"))
            if not abs_id:
                continue
            fmt, lvl_text = abstract_meta.get(abs_id, (None, None))
            if fmt == wanted:
                candidates.append((int(num_id_raw), abs_id, lvl_text))

        if not candidates:
            return None
        if kind == "ul":
            visible = [c for c in candidates if c[2] is not None and c[2].strip() != ""]
            if visible:
                return (visible[0][0], 0)
        return (candidates[0][0], 0)
    except Exception:
        return None


def fresh_numbering_for_kind(anchor: Paragraph, kind: str) -> tuple[int, int] | None:
    base = find_numbering_for_kind(anchor, kind)
    if base is None:
        return None
    abs_id = get_abstract_num_id_for_num_id(anchor, base[0])
    if abs_id is None:
        return None
    new_num_id = create_numbering_instance(anchor, abs_id, restart_at_one=(kind == "ol"))
    if new_num_id is None:
        return None
    return (new_num_id, 0)


INLINE_RE = re.compile(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)")
TABLE_ALIGN_RE = re.compile(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$")


def add_inline_runs(paragraph: Paragraph, text: str) -> None:
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def add_inline_runs_multiline(paragraph: Paragraph, lines: list[str]) -> None:
    for idx, line in enumerate(lines):
        if idx > 0:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        add_inline_runs(paragraph, line)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def split_table_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.clear()
    if not paragraph.runs:
        run = paragraph.add_run(text)
        run.bold = bold
    else:
        paragraph.runs[0].text = text
        paragraph.runs[0].bold = bold


def add_table(doc: Document, lines: list[str]) -> None:
    rows = [split_table_line(line) for line in lines]
    if len(rows) < 1:
        return
    data_rows = [rows[0]]
    for row in rows[1:]:
        if len(row) == 1 and TABLE_ALIGN_RE.match(lines[1].strip()):
            continue
        if all(set(cell) <= set("-: ") for cell in row):
            continue
        data_rows.append(row)
    col_count = max(len(row) for row in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=col_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(data_rows):
        for col_index in range(col_count):
            text = row[col_index] if col_index < len(row) else ""
            set_cell_text(table.cell(row_index, col_index), text, bold=(row_index == 0))
    doc.add_paragraph("")


def add_horizontal_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7B7B7")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)  # pylint: disable=protected-access


def convert_markdown(src: Path, dst: Path) -> None:
    doc = Document()
    configure_document(doc)

    lines = src.read_text(encoding="utf-8").splitlines()
    pending_blank = False
    index = 0
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if not stripped:
            pending_blank = True
            index += 1
            continue

        if is_table_line(raw):
            table_lines = []
            while index < len(lines) and is_table_line(lines[index]):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, table_lines)
            pending_blank = False
            continue

        if stripped == "---":
            add_horizontal_rule(doc)
            pending_blank = False
            index += 1
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph(style="Title")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            add_inline_runs(paragraph, stripped[2:].strip())
            pending_blank = False
            index += 1
            continue

        if stripped.startswith("## "):
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline_runs(paragraph, stripped[3:].strip())
            pending_blank = False
            index += 1
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline_runs(paragraph, stripped[4:].strip())
            pending_blank = False
            index += 1
            continue

        if stripped.startswith("> "):
            paragraph = doc.add_paragraph(style="Markdown Quote")
            add_inline_runs(paragraph, stripped[2:].strip())
            pending_blank = False
            index += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number_match:
            items = []
            while index < len(lines):
                current = lines[index].strip()
                current_match = re.match(r"^(\d+)\.\s+(.*)$", current)
                if not current_match:
                    break
                items.append(current_match.group(2).strip())
                index += 1
            anchor = doc.add_paragraph()
            anchor.paragraph_format.space_before = Pt(0)
            anchor.paragraph_format.space_after = Pt(0)
            numbering = fresh_numbering_for_kind(anchor, "ol") or find_numbering_for_kind(anchor, "ol")
            parent = anchor._element.getparent()
            if parent is not None:
                parent.remove(anchor._element)
            for item in items:
                paragraph = doc.add_paragraph(style="List Number")
                add_inline_runs(paragraph, item)
                if numbering is not None:
                    set_paragraph_numbering(paragraph, numbering[0], numbering[1])
            pending_blank = False
            continue

        if stripped.startswith("* ") or stripped.startswith("- "):
            items = []
            while index < len(lines):
                current = lines[index].strip()
                if not (current.startswith("* ") or current.startswith("- ")):
                    break
                items.append(current[2:].strip())
                index += 1
            for item in items:
                paragraph = doc.add_paragraph(style="List Bullet")
                add_inline_runs(paragraph, item)
            pending_blank = False
            continue

        paragraph_lines = [raw.strip()]
        index += 1
        while index < len(lines):
            next_raw = lines[index]
            next_stripped = next_raw.strip()
            if not next_stripped:
                break
            if (
                next_stripped.startswith("# ")
                or next_stripped.startswith("## ")
                or next_stripped.startswith("### ")
                or next_stripped.startswith("> ")
                or next_stripped.startswith("* ")
                or next_stripped.startswith("- ")
                or re.match(r"^(\d+)\.\s+", next_stripped)
                or next_stripped == "---"
                or is_table_line(lines[index])
            ):
                break
            paragraph_lines.append(next_raw.strip())
            index += 1

        paragraph = doc.add_paragraph()
        if pending_blank:
            paragraph.paragraph_format.space_before = Pt(4)
        add_inline_runs_multiline(paragraph, paragraph_lines)
        pending_blank = False

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Markdown zu DOCX mit kompaktem Standardsatz")
    parser.add_argument("--in", dest="infile", type=existing_markdown, required=True)
    parser.add_argument("--out", dest="outfile", type=Path, required=True)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    convert_markdown(args.infile, args.outfile)
    print(args.outfile)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(2)
