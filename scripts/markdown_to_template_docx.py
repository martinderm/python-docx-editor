#!/usr/bin/env python3
"""Convert Markdown to DOCX using a styled Word template.

Preserves template headers, footers, page size, margins, and styles.
Supports headings, lists, tables (styled), metadata lists (as tables),
and callout boxes (as single-cell tables).

Usage:
  python scripts/markdown_to_template_docx.py --in input.md --out output.docx [--template my_template.docx]
"""

import argparse
import sys
import re
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Regular expressions for inline formatting
INLINE_RE = re.compile(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|`[^`]+`|\*[^*]+\*)")

def add_inline_runs(paragraph, text: str) -> None:
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)

def set_cell_text(cell, text: str, bold: bool = False, header: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    for run in paragraph.runs:
        run.clear()
    
    # Check for markdown formatting in cell text
    add_inline_runs(paragraph, text)
    if bold:
        for run in paragraph.runs:
            run.bold = True
    if header:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0x00, 0x45, 0x7D)
            run.bold = True

def set_table_indentation(table, indent_dxa: int) -> None:
    if indent_dxa <= 0:
        return
    tblPr = table._tbl.tblPr
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is not None:
        tblPr.remove(tblInd)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(indent_dxa))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

def apply_paragraph_indent(p, default_left_indent, is_list: bool = False) -> None:
    if default_left_indent is not None:
        if is_list:
            p.paragraph_format.left_indent = default_left_indent + Pt(18)
        else:
            p.paragraph_format.left_indent = default_left_indent

def add_table(doc: Document, lines: list[str], indent_dxa: int = 0) -> None:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|") or not stripped.endswith("|"):
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        rows.append(cells)
        
    if not rows:
        return
        
    # Filter out alignment row (e.g. |---|---|)
    cleaned_rows = []
    for r in rows:
        if all(set(cell) <= set("-: ") for cell in r) and len(r) > 1:
            continue
        cleaned_rows.append(r)
        
    if not cleaned_rows:
        return
        
    col_count = max(len(r) for r in cleaned_rows)
    table = doc.add_table(rows=len(cleaned_rows), cols=col_count)
    table.style = "Table Grid"
    set_table_indentation(table, indent_dxa)
    
    # Make table look elegant
    for r_idx, row_data in enumerate(cleaned_rows):
        for c_idx in range(col_count):
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            is_header = (r_idx == 0)
            set_cell_text(table.cell(r_idx, c_idx), text, bold=is_header, header=is_header)

def add_metadata_table(doc: Document, items: list[tuple[str, str]], indent_dxa: int = 0) -> None:
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    set_table_indentation(table, indent_dxa)
    
    # Column 0: Key (bold, blue)
    # Column 1: Value (normal)
    for idx, (key, value) in enumerate(items):
        cell_key = table.cell(idx, 0)
        cell_val = table.cell(idx, 1)
        
        # Style key
        set_cell_text(cell_key, key.strip(), bold=True, header=True)
        # Style value
        set_cell_text(cell_val, value.strip(), bold=False, header=False)

def add_callout_box(doc: Document, text: str, indent_dxa: int = 0) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_indentation(table, indent_dxa)
    cell = table.cell(0, 0)
    set_cell_text(cell, text.strip(), bold=False, header=False)
    # Highlight block text
    for run in cell.paragraphs[0].runs:
        run.font.italic = True

def clear_doc_body(doc: Document) -> None:
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    for child in list(body):
        if child == sectPr:
            continue
        body.remove(child)


def ensure_required_styles(doc: Document) -> None:
    for style_name in ("List Bullet", "List Number"):
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
            style.font.name = "Aptos"
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(3)
    if "Table Grid" not in doc.styles:
        doc.styles.add_style("Table Grid", WD_STYLE_TYPE.TABLE)

def set_section_layout(section, source_layout, orientation) -> None:
    page_width, page_height, top_margin, bottom_margin, left_margin, right_margin = source_layout
    section.orientation = orientation
    if orientation == WD_ORIENT.LANDSCAPE:
        section.page_width = page_height
        section.page_height = page_width
    else:
        section.page_width = page_width
        section.page_height = page_height
    section.top_margin = top_margin
    section.bottom_margin = bottom_margin
    section.left_margin = left_margin
    section.right_margin = right_margin


def has_content_after(lines: list[str], index: int) -> bool:
    return any(line.strip() for line in lines[index:])


def convert_markdown_to_docx(
    md_path: Path,
    docx_path: Path,
    template_path: Path,
    table_orientation: str = "portrait",
) -> None:
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found at: {template_path}")
        
    doc = Document(str(template_path))
    ensure_required_styles(doc)
    
    # Read default left indents from template before clearing it
    default_left_indent = None
    default_table_indent_dxa = 0
    
    for p in doc.paragraphs:
        if p.paragraph_format.left_indent is not None:
            default_left_indent = p.paragraph_format.left_indent
            break
            
    for t in doc.tables:
        tblInd = t._tbl.tblPr.find(qn('w:tblInd'))
        if tblInd is not None:
            val = tblInd.get(qn('w:w'))
            if val:
                try:
                    default_table_indent_dxa = int(val)
                except ValueError:
                    pass
                break
                
    source_section = doc.sections[0]
    source_layout = (
        source_section.page_width,
        source_section.page_height,
        source_section.top_margin,
        source_section.bottom_margin,
        source_section.left_margin,
        source_section.right_margin,
    )
    clear_doc_body(doc)
    
    lines = md_path.read_text(encoding="utf-8").splitlines()
    
    index = 0
    in_metadata = False
    metadata_items = []
    
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        
        if not stripped:
            index += 1
            continue
            
        # Page breaks (Method 2: \pagebreak or \newpage)
        if stripped in {"\\pagebreak", "\\newpage"}:
            if metadata_items:
                add_metadata_table(doc, metadata_items, default_table_indent_dxa)
                metadata_items = []
                in_metadata = False
                doc.add_paragraph("")
            doc.add_page_break()
            index += 1
            continue
            
        # 1. Headings
        if stripped.startswith("# "):
            if metadata_items:
                add_metadata_table(doc, metadata_items, default_table_indent_dxa)
                metadata_items = []
                in_metadata = False
                doc.add_paragraph("")
                
            title_text = stripped[2:].strip()
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, title_text)
            doc.add_paragraph("") # Space after title
            index += 1
            continue
            
        if stripped.startswith("## "):
            if metadata_items:
                add_metadata_table(doc, metadata_items, default_table_indent_dxa)
                metadata_items = []
                doc.add_paragraph("")
                
            heading_text = stripped[3:].strip()
            if "metadata" in heading_text.lower() or "angaben" in heading_text.lower():
                in_metadata = True
            else:
                in_metadata = False
                
            h = doc.add_heading("", level=1)
            apply_paragraph_indent(h, default_left_indent)
            add_inline_runs(h, heading_text)
            index += 1
            continue
            
        if stripped.startswith("### "):
            if metadata_items:
                add_metadata_table(doc, metadata_items, default_table_indent_dxa)
                metadata_items = []
                doc.add_paragraph("")
            in_metadata = False
            heading_text = stripped[4:].strip()
            h = doc.add_heading("", level=2)
            apply_paragraph_indent(h, default_left_indent)
            add_inline_runs(h, heading_text)
            index += 1
            continue

        # 2. Blockquotes / Callout box
        if stripped.startswith("> "):
            blockquote_lines = []
            while index < len(lines) and lines[index].strip().startswith("> "):
                text_line = lines[index].strip()[2:].strip()
                if text_line.startswith("[!") and "]" in text_line:
                    text_line = text_line.split("]", 1)[1].strip()
                blockquote_lines.append(text_line)
                index += 1
            add_callout_box(doc, " ".join(blockquote_lines), default_table_indent_dxa)
            doc.add_paragraph("")
            continue
            
        # 3. Tables
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            if table_orientation == "landscape":
                set_section_layout(
                    doc.add_section(WD_SECTION.NEW_PAGE),
                    source_layout,
                    WD_ORIENT.LANDSCAPE,
                )
            add_table(doc, table_lines, default_table_indent_dxa)
            doc.add_paragraph("")
            if table_orientation == "landscape" and has_content_after(lines, index):
                set_section_layout(
                    doc.add_section(WD_SECTION.NEW_PAGE),
                    source_layout,
                    WD_ORIENT.PORTRAIT,
                )
            continue
            
        # 4. Lists (Bullet & Numbered)
        if stripped.startswith("* ") or stripped.startswith("- ") or re.match(r"^\d+\s*\.\s+", stripped):
            is_bullet = stripped.startswith("* ") or stripped.startswith("- ")
            list_items = []
            while index < len(lines):
                cur = lines[index].strip()
                if not cur:
                    index += 1
                    continue
                if is_bullet and (cur.startswith("* ") or cur.startswith("- ")):
                    list_items.append((cur[2:].strip(), "List Bullet"))
                    index += 1
                elif not is_bullet and re.match(r"^\d+\s*\.\s+", cur):
                    item_text = re.sub(r"^\d+\s*\.\s+", "", cur).strip()
                    list_items.append((item_text, "List Number"))
                    index += 1
                else:
                    break
            
            if in_metadata:
                for item_text, _ in list_items:
                    if ":" in item_text:
                        k, v = item_text.split(":", 1)
                        metadata_items.append((k.strip(), v.strip()))
                    else:
                        metadata_items.append((item_text.strip(), ""))
            else:
                for item_text, style_name in list_items:
                    p = doc.add_paragraph(style=style_name)
                    apply_paragraph_indent(p, default_left_indent, is_list=True)
                    add_inline_runs(p, item_text)
                doc.add_paragraph("")
            continue
            
        # 5. Normal paragraphs
        para_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line:
                break
            if (next_line.startswith("#") or 
                next_line.startswith("> ") or 
                next_line.startswith("|") or 
                next_line.startswith("* ") or 
                next_line.startswith("- ") or 
                re.match(r"^\d+\s*\.\s+", next_line)):
                break
            para_lines.append(next_line)
            index += 1
            
        p = doc.add_paragraph()
        apply_paragraph_indent(p, default_left_indent)
        add_inline_runs(p, " ".join(para_lines))
        doc.add_paragraph("")
        
    if metadata_items:
        add_metadata_table(doc, metadata_items, default_table_indent_dxa)
        doc.add_paragraph("")
        
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Successfully generated: {docx_path}")

def main() -> int:
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX using a styled Word template")
    parser.add_argument("--in", dest="infile", required=True, help="Input Markdown file")
    parser.add_argument("--out", dest="outfile", required=True, help="Output DOCX file")
    parser.add_argument("--template", dest="template", help="Word template file (.docx)")
    parser.add_argument(
        "--table-orientation",
        choices=("portrait", "landscape"),
        default="portrait",
        help="Page orientation for Markdown tables (default: portrait)",
    )
    args = parser.parse_args()
    
    infile = Path(args.infile)
    outfile = Path(args.outfile)
    
    if args.template:
        template = Path(args.template)
    else:
        # Default template path in python-docx-editor skill
        template = Path(__file__).parent.parent / "templates" / "GenericTemplate.docx"
        
    try:
        convert_markdown_to_docx(infile, outfile, template, args.table_orientation)
        return 0
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

if __name__ == "__main__":
    sys.exit(main())
