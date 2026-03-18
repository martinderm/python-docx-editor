#!/usr/bin/env python3
"""Project self-test for python-docx-editor.

Creates a temporary DOCX, runs v2 extraction and patch writeback,
and validates key contracts without committing test artifacts.

Usage:
  py -3 scripts/selftest.py
"""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
EXTRACT = ROOT / "scripts" / "extract_docx_for_llm.py"
PATCH = ROOT / "scripts" / "apply_docx_patch.py"
FILL_TABLE = ROOT / "scripts" / "fill_docx_table_from_json.py"


def run(cmd: list[str], cwd: Path) -> subprocess.CompletedProcess:
    return subprocess.run(cmd, cwd=str(cwd), text=True, capture_output=True, check=True)


def main() -> int:
    with tempfile.TemporaryDirectory(prefix="python-docx-editor-") as tmpdir:
        t = Path(tmpdir)

        in_docx = t / "in.docx"
        v2_json = t / "structure.v2.json"
        patch_json = t / "patch.json"
        out_docx = t / "out.docx"

        # Build minimal fixture document.
        d = Document()
        d.add_heading("1. Intro", level=1)
        d.add_paragraph("Hello base text")
        d.add_heading("2. Quality Principles", level=1)
        d.add_paragraph("Quality in MESHE is defined through five core principles:")
        d.add_paragraph("First item", style="List Bullet")
        d.add_paragraph("Second item", style="List Bullet")
        d.save(str(in_docx))

        # Extract v2 (default).
        run([sys.executable, str(EXTRACT), "--in", str(in_docx), "--out", str(v2_json)], ROOT)
        v2 = json.loads(v2_json.read_text(encoding="utf-8"))
        assert v2["schema"] == "docx-structure.v2", "Unexpected schema"

        # Find target paragraph block_id under section title.
        target_block = None
        for sec in v2["document"]["sections"]:
            if "Quality Principles" in sec.get("title", ""):
                for node in sec.get("content", []):
                    if node.get("type") == "paragraph" and "Quality in MESHE" in node.get("text", ""):
                        target_block = node.get("block_id")
                        break
        assert target_block, "Target block_id not found in v2 output"

        patch = {
            "ops": [
                {
                    "op": "replace_text",
                    "block_id": target_block,
                    "find": "Quality in MESHE is defined through five core principles:",
                    "replace": (
                        "Intro sentence for context. "
                        "Quality in MESHE is defined through five core principles:"
                    ),
                    "expected_matches": 1,
                }
            ]
        }
        patch_json.write_text(json.dumps(patch, ensure_ascii=False, indent=2), encoding="utf-8")

        # Apply patch.
        res = run(
            [
                sys.executable,
                str(PATCH),
                "--in",
                str(in_docx),
                "--out",
                str(out_docx),
                "--patch",
                str(patch_json),
            ],
            ROOT,
        )
        out = json.loads(res.stdout)
        assert out["results"][0]["status"] == "ok", "Patch operation did not succeed"

        d2 = Document(str(out_docx))
        texts = [" ".join(p.text.split()) for p in d2.paragraphs if p.text and p.text.strip()]
        assert any("Intro sentence for context." in t for t in texts), "Patched text missing"

        # Merge-aware table fill test (mixed merge patterns: 3+4 and 4+5).
        table_in = t / "table-in.docx"
        table_out = t / "table-out.docx"
        table_spec = t / "table-spec.json"

        td = Document()
        tt = td.add_table(rows=3, cols=6)
        for c in range(6):
            tt.cell(0, c).text = f"H{c + 1}"
        tt.rows[1].cells[3].merge(tt.rows[1].cells[4])
        tt.rows[2].cells[4].merge(tt.rows[2].cells[5])
        td.save(str(table_in))

        spec = {
            "table_index": 1,
            "rows": [
                {
                    "row_index": 2,
                    "mode": "fill",
                    "values": {
                        "output": "O1",
                        "risk": "R1",
                        "factors": "F1",
                        "mitigation": "M1",
                        "warning": "W1",
                    },
                },
                {
                    "row_index": 3,
                    "mode": "fill",
                    "values": {
                        "output": "O2",
                        "risk": "R2",
                        "factors": "F2",
                        "mitigation": "M2",
                        "warning": "W2",
                    },
                },
            ],
        }
        table_spec.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")

        run(
            [
                sys.executable,
                str(FILL_TABLE),
                "--in",
                str(table_in),
                "--out",
                str(table_out),
                "--spec",
                str(table_spec),
            ],
            ROOT,
        )

        td2 = Document(str(table_out))
        tr1 = td2.tables[0].rows[1]
        tr2 = td2.tables[0].rows[2]
        assert tr1.cells[3].text == "M1" and tr1.cells[5].text == "W1", "merge_3_4 mapping failed"
        assert tr2.cells[3].text == "M2" and tr2.cells[4].text == "W2", "merge_4_5 mapping failed"

    print("selftest: OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
