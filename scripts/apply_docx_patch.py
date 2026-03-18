#!/usr/bin/env python3
"""Apply minimal, block-targeted patches to .docx files.

Designed to work with block_ids produced by extract_docx_for_llm.py.

Supports plain text ops and markdown-aware rewrite ops.
"""

from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
import sys

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class ReplaceResult:
    matches: int = 0
    changes: int = 0
    cross_run_conflicts: int = 0


def existing_docx(path_str: str) -> Path:
    p = Path(path_str)
    if not p.exists():
        raise argparse.ArgumentTypeError(f"Datei nicht gefunden: {path_str}")
    if p.suffix.lower() != ".docx":
        raise argparse.ArgumentTypeError(f"Keine .docx-Datei: {path_str}")
    return p


def existing_json(path_str: str) -> Path:
    p = Path(path_str)
    if not p.exists():
        raise argparse.ArgumentTypeError(f"Patch-Datei nicht gefunden: {path_str}")
    if p.suffix.lower() not in {".json", ".jsonl"}:
        raise argparse.ArgumentTypeError(f"Keine JSON-Datei: {path_str}")
    return p


def iter_block_items(parent: DocumentObject):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def parse_table_block_id(block_id: str) -> tuple[int, int, int] | None:
    m = re.fullmatch(r"t_(\d+)_r(\d+)_(\d+)", block_id)
    if not m:
        return None
    return int(m.group(1)), int(m.group(2)), int(m.group(3))


def replace_in_runs(paragraph: Paragraph, old: str, new: str) -> ReplaceResult:
    res = ReplaceResult()
    if not old:
        return res

    para_text = paragraph.text
    if old not in para_text:
        return res

    res.matches = para_text.count(old)

    for run in paragraph.runs:
        if old in run.text:
            c = run.text.count(old)
            run.text = run.text.replace(old, new)
            res.changes += c

    if res.changes < res.matches:
        res.cross_run_conflicts = res.matches - res.changes

    return res


def build_paragraph_index(doc: DocumentObject) -> dict[str, Paragraph]:
    out: dict[str, Paragraph] = {}
    p_i = 0
    h_i = 0
    f_i = 0

    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            if not item.text or not item.text.strip():
                continue
            p_i += 1
            out[f"p_{p_i}"] = item

    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                h_i += 1
                out[f"h_{h_i}"] = p
        for p in section.footer.paragraphs:
            if p.text.strip():
                f_i += 1
                out[f"f_{f_i}"] = p

    return out


def build_table_index(doc: DocumentObject) -> dict[int, Table]:
    out: dict[int, Table] = {}
    t_i = 0
    for item in iter_block_items(doc):
        if isinstance(item, Table):
            t_i += 1
            out[t_i] = item
    return out


def replace_in_table_range(table: Table, row_start: int, row_end: int, old: str, new: str) -> ReplaceResult:
    res = ReplaceResult()
    if not old:
        return res

    for row_idx in range(max(1, row_start), min(len(table.rows), row_end) + 1):
        row = table.rows[row_idx - 1]
        for cell in row.cells:
            for p in cell.paragraphs:
                p_res = replace_in_runs(p, old, new)
                res.matches += p_res.matches
                res.changes += p_res.changes
                res.cross_run_conflicts += p_res.cross_run_conflicts

    return res


def load_patch(path: Path) -> dict:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict) or "ops" not in data or not isinstance(data["ops"], list):
        raise ValueError("Patch-Datei muss ein JSON-Objekt mit 'ops' (Liste) sein.")
    return data


def clear_paragraph_numbering(paragraph: Paragraph) -> None:
    ppr = paragraph._p.pPr  # pylint: disable=protected-access
    if ppr is not None and ppr.numPr is not None:
        ppr.remove(ppr.numPr)


def get_paragraph_numbering(paragraph: Paragraph) -> tuple[int, int] | None:
    ppr = paragraph._p.pPr  # pylint: disable=protected-access
    if ppr is None or ppr.numPr is None:
        return None
    num_pr = ppr.numPr
    num_id_el = num_pr.find(qn("w:numId"))
    ilvl_el = num_pr.find(qn("w:ilvl"))
    if num_id_el is None:
        return None
    try:
        num_id = int(num_id_el.get(qn("w:val")))
        ilvl = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0
    except Exception:
        return None
    return (num_id, ilvl)


def set_paragraph_numbering(paragraph: Paragraph, num_id: int, ilvl: int = 0) -> None:
    ppr = paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
    if ppr.numPr is not None:
        ppr.remove(ppr.numPr)

    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    ppr.append(num_pr)


def get_abstract_num_id_for_num_id(anchor: Paragraph, num_id: int) -> int | None:
    try:
        root = anchor.part.numbering_part._element  # pylint: disable=protected-access
        for num in root.findall(qn("w:num")):
            num_id_raw = num.get(qn("w:numId"))
            if not num_id_raw or int(num_id_raw) != int(num_id):
                continue
            abs_ref = num.find(qn("w:abstractNumId"))
            if abs_ref is None:
                return None
            abs_id_raw = abs_ref.get(qn("w:val"))
            return int(abs_id_raw) if abs_id_raw is not None else None
    except Exception:
        return None
    return None


def numbering_format_for_abstract_num(anchor: Paragraph, abstract_num_id: int) -> str | None:
    try:
        root = anchor.part.numbering_part._element  # pylint: disable=protected-access
        for abstract in root.findall(qn("w:abstractNum")):
            abs_id_raw = abstract.get(qn("w:abstractNumId"))
            if not abs_id_raw or int(abs_id_raw) != int(abstract_num_id):
                continue
            lvl0 = None
            for lvl in abstract.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl"))
                if ilvl in {"0", None}:
                    lvl0 = lvl
                    break
            if lvl0 is None:
                return None
            num_fmt_el = lvl0.find(qn("w:numFmt"))
            if num_fmt_el is None:
                return None
            return num_fmt_el.get(qn("w:val"))
    except Exception:
        return None
    return None


def create_numbering_instance(anchor: Paragraph, abstract_num_id: int, *, restart_at_one: bool = False) -> int | None:
    """Create a fresh w:num for a given abstractNum, returning new numId.

    restart_at_one adds lvlOverride/startOverride=1 for ilvl=0.
    """
    try:
        root = anchor.part.numbering_part._element  # pylint: disable=protected-access
        max_num_id = 0
        for num in root.findall(qn("w:num")):
            raw = num.get(qn("w:numId"))
            if not raw:
                continue
            max_num_id = max(max_num_id, int(raw))

        new_num_id = max_num_id + 1
        num_el = OxmlElement("w:num")
        num_el.set(qn("w:numId"), str(new_num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(int(abstract_num_id)))
        num_el.append(abs_ref)

        if restart_at_one:
            lvl_override = OxmlElement("w:lvlOverride")
            lvl_override.set(qn("w:ilvl"), "0")
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), "1")
            lvl_override.append(start_override)
            num_el.append(lvl_override)

        root.append(num_el)
        return new_num_id
    except Exception:
        return None


def style_looks_like_list(style_name: str | None) -> bool:
    if not style_name:
        return False
    s = style_name.lower()
    return any(k in s for k in ["list", "bullet", "number", "aufz", "aufl"])


def clear_runs(paragraph: Paragraph) -> None:
    for r in list(paragraph._p.r_lst):  # pylint: disable=protected-access
        paragraph._p.remove(r)  # pylint: disable=protected-access


def add_hyperlink(paragraph: Paragraph, text: str, url: str, *, bold: bool = False, italic: bool = False) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    if italic:
        i = OxmlElement("w:i")
        r_pr.append(i)

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)  # pylint: disable=protected-access


def _unescape_md(text: str) -> str:
    return re.sub(r"\\([`*_\[\]()|>#-])", r"\1", text)


def render_inline_markdown(paragraph: Paragraph, text: str) -> None:
    clear_runs(paragraph)
    rest = text

    # order matters: link, bold+italic, bold, italic, code
    patterns = [
        ("link", re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")),
        ("bolditalic", re.compile(r"\*\*\*(.+?)\*\*\*")),
        ("bold", re.compile(r"\*\*(.+?)\*\*")),
        ("italic", re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")),
        ("code", re.compile(r"`([^`]+)`")),
    ]

    while rest:
        earliest = None
        earliest_kind = None
        earliest_match = None
        for kind, rx in patterns:
            m = rx.search(rest)
            if m and (earliest is None or m.start() < earliest):
                earliest = m.start()
                earliest_kind = kind
                earliest_match = m

        if earliest_match is None:
            paragraph.add_run(_unescape_md(rest))
            break

        if earliest_match.start() > 0:
            paragraph.add_run(_unescape_md(rest[: earliest_match.start()]))

        if earliest_kind == "link":
            add_hyperlink(paragraph, earliest_match.group(1), earliest_match.group(2))
        elif earliest_kind == "bolditalic":
            run = paragraph.add_run(_unescape_md(earliest_match.group(1)))
            run.bold = True
            run.italic = True
        elif earliest_kind == "bold":
            run = paragraph.add_run(_unescape_md(earliest_match.group(1)))
            run.bold = True
        elif earliest_kind == "italic":
            run = paragraph.add_run(_unescape_md(earliest_match.group(1)))
            run.italic = True
        elif earliest_kind == "code":
            run = paragraph.add_run(_unescape_md(earliest_match.group(1)))
            run.font.name = "Consolas"

        rest = rest[earliest_match.end() :]


def apply_text_or_runs(
    paragraph: Paragraph,
    *,
    text: str | None = None,
    runs: list[dict] | None = None,
    markdown: bool = False,
) -> None:
    if runs is not None:
        clear_runs(paragraph)
        for r in runs:
            if not isinstance(r, dict):
                raise ValueError("runs[] enthält Nicht-Objekt.")
            r_text = r.get("text", "")
            if not isinstance(r_text, str):
                raise ValueError("runs[].text muss String sein.")
            link = r.get("link")
            bold = bool(r.get("bold", False))
            italic = bool(r.get("italic", False))
            code = bool(r.get("code", False))
            if link:
                if not isinstance(link, str):
                    raise ValueError("runs[].link muss String (URL) sein.")
                add_hyperlink(paragraph, r_text, link, bold=bold, italic=italic)
            else:
                run = paragraph.add_run(r_text)
                run.bold = bold
                run.italic = italic
                if code:
                    run.font.name = "Consolas"
        return

    if text is None:
        raise ValueError("Es muss entweder text oder runs gesetzt sein.")

    if markdown:
        render_inline_markdown(paragraph, text)
    else:
        paragraph.text = text


def apply_replace_op(op: dict, paragraph_index: dict[str, Paragraph], table_index: dict[int, Table]) -> dict:
    block_id = op.get("block_id")
    old = op.get("find")
    new = op.get("replace", "")
    expected = int(op.get("expected_matches", 1))

    if not isinstance(block_id, str) or not block_id:
        raise ValueError("replace_text op braucht 'block_id'.")
    if not isinstance(old, str) or old == "":
        raise ValueError("replace_text op braucht nicht-leeres 'find'.")
    if not isinstance(new, str):
        raise ValueError("replace_text op braucht String in 'replace'.")

    table_spec = parse_table_block_id(block_id)
    if table_spec is not None:
        t_idx, r_start, r_end = table_spec
        table = table_index.get(t_idx)
        if table is None:
            raise ValueError(f"Table block_id nicht gefunden: {block_id}")
        res = replace_in_table_range(table, r_start, r_end, old, new)
    else:
        paragraph = paragraph_index.get(block_id)
        if paragraph is None:
            raise ValueError(f"Paragraph/Header/Footer block_id nicht gefunden: {block_id}")
        res = replace_in_runs(paragraph, old, new)

    if res.matches != expected:
        raise ValueError(
            f"Op auf {block_id}: expected_matches={expected}, gefunden={res.matches}. "
            "Keine stillen Mehrfachtreffer erlaubt."
        )
    if res.cross_run_conflicts > 0:
        raise ValueError(
            f"Op auf {block_id}: {res.cross_run_conflicts} Treffer über Run-Grenzen. "
            "Minimal-Edit verweigert (kein Full-Rewrite)."
        )

    return {
        "op": "replace_text",
        "block_id": block_id,
        "matches": res.matches,
        "changes": res.changes,
        "status": "ok",
    }


def apply_set_paragraph_op(op: dict, paragraph_index: dict[str, Paragraph]) -> dict:
    block_id = op.get("block_id")
    text = op.get("text")
    runs = op.get("runs")
    style = op.get("style")
    markdown = bool(op.get("markdown", False))
    expected_contains = op.get("expected_contains")
    clear_list_format = bool(op.get("clear_list_format", True))

    if not isinstance(block_id, str) or not block_id:
        raise ValueError("set_paragraph op braucht 'block_id'.")
    if runs is None and not isinstance(text, str):
        raise ValueError("set_paragraph op braucht String in 'text' oder runs[].")
    if style is not None and not isinstance(style, str):
        raise ValueError("set_paragraph op: 'style' muss String sein, wenn gesetzt.")
    if expected_contains is not None and not isinstance(expected_contains, str):
        raise ValueError("set_paragraph op: 'expected_contains' muss String sein, wenn gesetzt.")

    if parse_table_block_id(block_id) is not None:
        raise ValueError("set_paragraph unterstützt keine table block_ids (t_*).")

    paragraph = paragraph_index.get(block_id)
    if paragraph is None:
        raise ValueError(f"Paragraph/Header/Footer block_id nicht gefunden: {block_id}")

    current = paragraph.text or ""
    if expected_contains is not None and expected_contains not in current:
        raise ValueError(
            f"Op auf {block_id}: expected_contains nicht gefunden. "
            "Abbruch, um falsche Zielstellen zu vermeiden."
        )

    if style:
        paragraph.style = style
    apply_text_or_runs(paragraph, text=text if isinstance(text, str) else None, runs=runs, markdown=markdown)

    if clear_list_format and not style_looks_like_list(style):
        clear_paragraph_numbering(paragraph)

    return {
        "op": "set_paragraph",
        "block_id": block_id,
        "status": "ok",
        "style": style if style else None,
        "new_length": len(paragraph.text or ""),
    }


def apply_delete_paragraph_op(op: dict, paragraph_index: dict[str, Paragraph]) -> dict:
    block_id = op.get("block_id")
    expected_contains = op.get("expected_contains")

    if not isinstance(block_id, str) or not block_id:
        raise ValueError("delete_paragraph op braucht 'block_id'.")
    if expected_contains is not None and not isinstance(expected_contains, str):
        raise ValueError("delete_paragraph op: 'expected_contains' muss String sein, wenn gesetzt.")
    if parse_table_block_id(block_id) is not None:
        raise ValueError("delete_paragraph unterstützt keine table block_ids (t_*).")

    paragraph = paragraph_index.get(block_id)
    if paragraph is None:
        raise ValueError(f"Paragraph/Header/Footer block_id nicht gefunden: {block_id}")

    current = paragraph.text or ""
    if expected_contains is not None and expected_contains not in current:
        raise ValueError(
            f"Op auf {block_id}: expected_contains nicht gefunden. "
            "Abbruch, um falsche Zielstellen zu vermeiden."
        )

    p = paragraph._element  # pylint: disable=protected-access
    parent = p.getparent()
    if parent is None:
        raise ValueError(f"Op auf {block_id}: Paragraph hat kein Parent-Element.")
    parent.remove(p)

    return {"op": "delete_paragraph", "block_id": block_id, "status": "ok"}


def build_body_paragraph_order(doc: DocumentObject) -> list[str]:
    order: list[str] = []
    p_i = 0
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            if not item.text or not item.text.strip():
                continue
            p_i += 1
            order.append(f"p_{p_i}")
    return order


def insert_paragraph_after(paragraph: Paragraph, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)  # pylint: disable=protected-access
    new_para = Paragraph(new_p, paragraph._parent)  # pylint: disable=protected-access
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def set_horizontal_rule(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)


def insert_table_after(paragraph: Paragraph, rows: list[list[str]], table_style: str | None = None) -> Paragraph:
    max_cols = max((len(r) for r in rows), default=1)
    body = paragraph._parent  # pylint: disable=protected-access
    table = body.add_table(rows=len(rows), cols=max_cols)
    if table_style:
        try:
            table.style = table_style
        except Exception:
            pass
    for r_i, row in enumerate(rows):
        for c_i in range(max_cols):
            table.cell(r_i, c_i).text = row[c_i] if c_i < len(row) else ""

    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)  # pylint: disable=protected-access

    after = OxmlElement("w:p")
    tbl.addnext(after)
    return Paragraph(after, paragraph._parent)  # pylint: disable=protected-access


def parse_markdown_blocks(markdown: str) -> list[dict]:
    lines = markdown.splitlines()
    i = 0
    blocks: list[dict] = []

    def flush_para(buf: list[str]) -> None:
        if buf:
            blocks.append({"type": "paragraph", "text": " ".join(s.strip() for s in buf if s.strip())})
            buf.clear()

    para_buf: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_para(para_buf)
            i += 1
            continue

        # table
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?[\s:\-\|]+\|?\s*$", lines[i + 1].strip()):
            flush_para(para_buf)
            table_lines = [stripped]
            i += 2
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for tl in table_lines:
                t = tl.strip().strip("|")
                cells = [c.strip() for c in t.split("|")]
                rows.append(cells)
            blocks.append({"type": "table", "rows": rows})
            continue

        # heading
        hm = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if hm:
            flush_para(para_buf)
            blocks.append({"type": "heading", "level": len(hm.group(1)), "text": hm.group(2).strip()})
            i += 1
            continue

        # horizontal rule
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", stripped):
            flush_para(para_buf)
            blocks.append({"type": "hr"})
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            flush_para(para_buf)
            qbuf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                qbuf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append({"type": "quote", "text": " ".join(qbuf)})
            continue

        # unordered list
        um = re.match(r"^\s*[-*]\s+(.+)$", line)
        if um:
            flush_para(para_buf)
            items = []
            while i < len(lines):
                m = re.match(r"^\s*[-*]\s+(.+)$", lines[i])
                if not m:
                    break
                items.append(m.group(1).strip())
                i += 1
            blocks.append({"type": "ul", "items": items})
            continue

        # ordered list
        om = re.match(r"^\s*\d+[\.)]\s+(.+)$", line)
        if om:
            flush_para(para_buf)
            items = []
            while i < len(lines):
                m = re.match(r"^\s*\d+[\.)]\s+(.+)$", lines[i])
                if not m:
                    break
                items.append(m.group(1).strip())
                i += 1
            blocks.append({"type": "ol", "items": items})
            continue

        para_buf.append(line)
        i += 1

    flush_para(para_buf)
    return blocks


def find_numbering_for_kind(anchor: Paragraph, kind: str) -> tuple[int, int] | None:
    """Find an existing numbering definition in the document for bullet/decimal lists.

    Returns (num_id, ilvl=0) when possible, otherwise None.
    For bullets, prefers abstract definitions with a visible lvlText symbol.
    """
    wanted = "bullet" if kind == "ul" else "decimal"
    try:
        root = anchor.part.numbering_part._element  # pylint: disable=protected-access

        # Map abstractNumId -> (lvl0 numFmt, lvl0 lvlText)
        abstract_meta: dict[str, tuple[str | None, str | None]] = {}
        for abstract in root.findall(qn("w:abstractNum")):
            abs_id = abstract.get(qn("w:abstractNumId"))
            if not abs_id:
                continue
            lvl0 = None
            for lvl in abstract.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl"))
                if ilvl in {"0", None}:
                    lvl0 = lvl
                    break
            if lvl0 is None:
                continue
            num_fmt_el = lvl0.find(qn("w:numFmt"))
            num_fmt_val = num_fmt_el.get(qn("w:val")) if num_fmt_el is not None else None
            lvl_text_el = lvl0.find(qn("w:lvlText"))
            lvl_text_val = lvl_text_el.get(qn("w:val")) if lvl_text_el is not None else None
            abstract_meta[abs_id] = (num_fmt_val, lvl_text_val)

        candidates: list[tuple[int, str, str | None]] = []
        for num in root.findall(qn("w:num")):
            num_id_raw = num.get(qn("w:numId"))
            abs_ref = num.find(qn("w:abstractNumId"))
            if not num_id_raw or abs_ref is None:
                continue
            abs_id = abs_ref.get(qn("w:val"))
            if not abs_id:
                continue
            fmt, lvl_text = abstract_meta.get(abs_id, (None, None))
            if fmt == wanted:
                candidates.append((int(num_id_raw), abs_id, lvl_text))

        if not candidates:
            return None

        if kind == "ul":
            # Prefer visible bullet glyph definitions; some templates have blank bullet lvlText.
            visible = [c for c in candidates if (c[2] is not None and c[2].strip() != "")]
            if visible:
                return (visible[0][0], 0)

        return (candidates[0][0], 0)
    except Exception:
        return None

    return None


def numbering_kind_for_num_id(anchor: Paragraph, num_id: int) -> str | None:
    abs_id = get_abstract_num_id_for_num_id(anchor, num_id)
    if abs_id is None:
        return None
    return numbering_format_for_abstract_num(anchor, abs_id)


def fresh_numbering_for_kind(anchor: Paragraph, kind: str) -> tuple[int, int] | None:
    """Return a fresh numbering instance (new numId) for ul/ol kind.

    For ordered lists, force explicit restart at 1.
    """
    base = find_numbering_for_kind(anchor, kind)
    if base is None:
        return None
    abs_id = get_abstract_num_id_for_num_id(anchor, base[0])
    if abs_id is None:
        return None
    new_num_id = create_numbering_instance(anchor, abs_id, restart_at_one=(kind == "ol"))
    if new_num_id is None:
        return None
    return (new_num_id, 0)


def render_markdown_blocks_after(
    anchor: Paragraph,
    markdown: str,
    ul_numbering: tuple[int, int] | None = None,
    ol_numbering: tuple[int, int] | None = None,
    list_mode: str = "auto",
) -> tuple[Paragraph, int, dict]:
    blocks = parse_markdown_blocks(markdown)
    inserted = 0
    cur = anchor
    list_stats = {
        "list_mode": list_mode,
        "ooxml_fallback_used": False,
        "ol_restart_forced": False,
        "ul_fallback_reason": None,
    }

    for b in blocks:
        btype = b.get("type")
        if btype == "table":
            cur = insert_table_after(cur, b.get("rows", []), table_style="Table Grid")
            inserted += 1
            continue

        if btype == "hr":
            p = insert_paragraph_after(cur)
            set_horizontal_rule(p)
            cur = p
            inserted += 1
            continue

        if btype == "heading":
            level = int(b.get("level", 1))
            style = f"Heading {max(1, min(6, level))}"
            p = insert_paragraph_after(cur, style=style)
            render_inline_markdown(p, b.get("text", ""))
            clear_paragraph_numbering(p)
            cur = p
            inserted += 1
            continue

        if btype == "quote":
            p = insert_paragraph_after(cur, style="Quote")
            render_inline_markdown(p, b.get("text", ""))
            clear_paragraph_numbering(p)
            cur = p
            inserted += 1
            continue

        if btype in {"ul", "ol"}:
            style = "List Bullet" if btype == "ul" else "List Number"
            numbering = ul_numbering if btype == "ul" else ol_numbering

            # Validate hint kind (bullet vs decimal). Mismatched hints are ignored.
            if numbering is not None:
                hinted_kind = numbering_kind_for_num_id(anchor, numbering[0])
                wanted_kind = "bullet" if btype == "ul" else "decimal"
                if hinted_kind != wanted_kind:
                    numbering = None

            # No hint -> auto-detect by kind from template.
            if numbering is None:
                numbering = find_numbering_for_kind(anchor, btype)

            use_numpr = list_mode == "ooxml"
            if list_mode == "auto" and btype == "ol":
                # For stable chapter-local numbering, force restart for ordered lists.
                use_numpr = True
                list_stats["ol_restart_forced"] = True

            # Ordered lists: if numPr is used, create fresh numbering instance per block.
            if btype == "ol" and use_numpr:
                fresh = fresh_numbering_for_kind(anchor, "ol")
                if fresh is not None:
                    numbering = fresh
                    list_stats["ooxml_fallback_used"] = True

            for item in b.get("items", []):
                p = insert_paragraph_after(cur, style=style)
                render_inline_markdown(p, item)

                style_name = p.style.name if p.style is not None else None
                style_ok = style_looks_like_list(style_name)

                if btype == "ul" and list_mode == "auto" and not style_ok:
                    # Style-only failed (template/style mapping issue) -> fallback to OOXML bullets.
                    if numbering is None:
                        numbering = fresh_numbering_for_kind(anchor, "ul") or find_numbering_for_kind(anchor, "ul")
                    use_numpr = numbering is not None
                    list_stats["ooxml_fallback_used"] = use_numpr
                    if use_numpr and list_stats["ul_fallback_reason"] is None:
                        list_stats["ul_fallback_reason"] = "style_not_list_like"

                if use_numpr and numbering is not None:
                    set_paragraph_numbering(p, numbering[0], numbering[1])

                cur = p
                inserted += 1
            continue

        # paragraph fallback
        p = insert_paragraph_after(cur, style="Normal")
        render_inline_markdown(p, b.get("text", ""))
        clear_paragraph_numbering(p)
        cur = p
        inserted += 1

    return cur, inserted, list_stats


def apply_replace_paragraph_range_op(op: dict, doc: DocumentObject, paragraph_index: dict[str, Paragraph]) -> dict:
    start_block_id = op.get("start_block_id")
    end_block_id = op.get("end_block_id")
    new_paragraphs = op.get("new_paragraphs")
    expected_start_contains = op.get("expected_start_contains")
    expected_end_contains = op.get("expected_end_contains")
    allow_headings = bool(op.get("allow_headings", False))

    if not isinstance(start_block_id, str) or not start_block_id.startswith("p_"):
        raise ValueError("replace_paragraph_range braucht 'start_block_id' vom Typ p_<n>.")
    if not isinstance(end_block_id, str) or not end_block_id.startswith("p_"):
        raise ValueError("replace_paragraph_range braucht 'end_block_id' vom Typ p_<n>.")
    if not isinstance(new_paragraphs, list) or len(new_paragraphs) == 0:
        raise ValueError("replace_paragraph_range braucht nicht-leere 'new_paragraphs'-Liste.")

    start_para = paragraph_index.get(start_block_id)
    end_para = paragraph_index.get(end_block_id)
    if start_para is None or end_para is None:
        raise ValueError("replace_paragraph_range: start/end block_id nicht gefunden.")

    if expected_start_contains is not None:
        if not isinstance(expected_start_contains, str):
            raise ValueError("expected_start_contains muss String sein.")
        if expected_start_contains not in (start_para.text or ""):
            raise ValueError("replace_paragraph_range: expected_start_contains nicht gefunden.")

    if expected_end_contains is not None:
        if not isinstance(expected_end_contains, str):
            raise ValueError("expected_end_contains muss String sein.")
        if expected_end_contains not in (end_para.text or ""):
            raise ValueError("replace_paragraph_range: expected_end_contains nicht gefunden.")

    order = build_body_paragraph_order(doc)
    try:
        i_start = order.index(start_block_id)
        i_end = order.index(end_block_id)
    except ValueError as exc:
        raise ValueError("replace_paragraph_range: start/end block_id nicht in Body-Reihenfolge gefunden.") from exc
    if i_start > i_end:
        raise ValueError("replace_paragraph_range: start_block_id muss vor end_block_id liegen.")

    old_ids = order[i_start : i_end + 1]
    old_paragraphs = [paragraph_index[bid] for bid in old_ids]

    if not allow_headings:
        heading_hits = []
        for bid, p in zip(old_ids, old_paragraphs):
            style_name = p.style.name if p.style else ""
            if style_name and (style_name.startswith("Heading") or style_name.startswith("Überschrift")):
                heading_hits.append(f"{bid}({style_name})")
        if heading_hits:
            raise ValueError(
                "replace_paragraph_range enthält Heading-Absätze: "
                + ", ".join(heading_hits)
                + ". Setze allow_headings=true nur wenn das absichtlich ist."
            )

    anchor = end_para
    inserted = 0
    for entry in new_paragraphs:
        if not isinstance(entry, dict):
            raise ValueError("replace_paragraph_range: jedes Element in new_paragraphs muss Objekt sein.")

        style = entry.get("style")
        text = entry.get("text")
        runs = entry.get("runs")
        markdown = bool(entry.get("markdown", False))

        if style is not None and not isinstance(style, str):
            raise ValueError("replace_paragraph_range: new_paragraphs[].style muss String sein, wenn gesetzt.")
        if runs is None and not isinstance(text, str):
            raise ValueError("replace_paragraph_range: new_paragraphs[] braucht text oder runs[].")

        p = insert_paragraph_after(anchor, style=style)
        apply_text_or_runs(p, text=text if isinstance(text, str) else None, runs=runs, markdown=markdown)
        if not style_looks_like_list(style):
            clear_paragraph_numbering(p)
        anchor = p
        inserted += 1

    removed_preview = [" ".join((p.text or "").split())[:120] for p in old_paragraphs]

    for p in old_paragraphs:
        el = p._element  # pylint: disable=protected-access
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    return {
        "op": "replace_paragraph_range",
        "start_block_id": start_block_id,
        "end_block_id": end_block_id,
        "removed": len(old_paragraphs),
        "inserted": inserted,
        "removed_preview": removed_preview,
        "status": "ok",
    }


def apply_replace_paragraph_range_markdown_op(op: dict, doc: DocumentObject, paragraph_index: dict[str, Paragraph]) -> dict:
    start_block_id = op.get("start_block_id")
    end_block_id = op.get("end_block_id")
    markdown = op.get("markdown")
    allow_headings = bool(op.get("allow_headings", False))

    if not isinstance(markdown, str) or not markdown.strip():
        raise ValueError("replace_paragraph_range_markdown braucht nicht-leeren markdown-String.")

    # Best-effort numbering hint from replaced range (for templates without List Bullet/List Number styles).
    ul_hint = None
    ol_hint = None
    try:
        order = build_body_paragraph_order(doc)
        i_start = order.index(start_block_id)
        i_end = order.index(end_block_id)
        if i_start <= i_end:
            old_ids = order[i_start : i_end + 1]
            old_paragraphs = [paragraph_index[bid] for bid in old_ids if bid in paragraph_index]
            old_numberings = [n for n in (get_paragraph_numbering(p) for p in old_paragraphs) if n is not None]
            if old_numberings:
                md_blocks = parse_markdown_blocks(markdown)
                has_ul = any(b.get("type") == "ul" for b in md_blocks)
                has_ol = any(b.get("type") == "ol" for b in md_blocks)
                if has_ul and not has_ol:
                    ul_hint = old_numberings[0]
                elif has_ol and not has_ul:
                    ol_hint = old_numberings[0]
                elif has_ul and has_ol:
                    ul_hint = old_numberings[0]
                    ol_hint = old_numberings[1] if len(old_numberings) > 1 else old_numberings[0]
    except Exception:
        # Hinting is optional; on any issue we continue without numbering hints.
        ul_hint = None
        ol_hint = None

    # Reuse validation and range selection by calling base op with placeholder paragraph.
    base_result = apply_replace_paragraph_range_op(
        {
            "start_block_id": start_block_id,
            "end_block_id": end_block_id,
            "new_paragraphs": [{"text": "__TMP__", "style": "Normal"}],
            "allow_headings": allow_headings,
            "expected_start_contains": op.get("expected_start_contains"),
            "expected_end_contains": op.get("expected_end_contains"),
        },
        doc,
        paragraph_index,
    )

    # remove tmp paragraph and insert markdown blocks in its place
    p_index = build_paragraph_index(doc)
    # find temp paragraph by marker
    tmp_para = next((p for p in p_index.values() if (p.text or "") == "__TMP__"), None)
    if tmp_para is None:
        raise ValueError("replace_paragraph_range_markdown: temporärer Absatz nicht gefunden.")

    anchor = tmp_para
    parent = tmp_para._element.getparent()  # pylint: disable=protected-access
    if parent is None:
        raise ValueError("replace_paragraph_range_markdown: kein Parent für temp Absatz.")

    # delete tmp first, keep previous sibling as anchor base
    prev = tmp_para._element.getprevious()  # pylint: disable=protected-access
    parent.remove(tmp_para._element)  # pylint: disable=protected-access
    if prev is None:
        # create anchor paragraph at top if needed
        new_p = OxmlElement("w:p")
        body = doc.element.body
        body.insert(0, new_p)
        anchor = Paragraph(new_p, doc)
    else:
        anchor = Paragraph(prev, doc)

    _, inserted_blocks, list_stats = render_markdown_blocks_after(
        anchor,
        markdown,
        ul_numbering=ul_hint,
        ol_numbering=ol_hint,
        list_mode="auto",
    )
    base_result["op"] = "replace_paragraph_range_markdown"
    base_result["inserted_markdown_blocks"] = inserted_blocks
    base_result["list_handling"] = list_stats
    if ul_hint is not None:
        base_result["ul_numbering_hint"] = {"num_id": ul_hint[0], "level": ul_hint[1]}
    if ol_hint is not None:
        base_result["ol_numbering_hint"] = {"num_id": ol_hint[0], "level": ol_hint[1]}
    return base_result


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Apply minimal block-targeted DOCX patches")
    p.add_argument("--in", dest="infile", type=existing_docx, required=True)
    p.add_argument("--out", dest="outfile", type=Path, required=True)
    p.add_argument("--patch", dest="patchfile", type=existing_json, required=True)
    return p.parse_args()


def main() -> int:
    args = parse_args()
    patch = load_patch(args.patchfile)

    doc = Document(str(args.infile))

    results: list[dict] = []
    for i, op in enumerate(patch["ops"], start=1):
        if not isinstance(op, dict):
            raise ValueError(f"Op #{i} ist kein Objekt.")

        paragraph_index = build_paragraph_index(doc)
        table_index = build_table_index(doc)

        op_kind = op.get("op")
        if op_kind == "replace_text":
            results.append(apply_replace_op(op, paragraph_index, table_index))
        elif op_kind == "set_paragraph":
            results.append(apply_set_paragraph_op(op, paragraph_index))
        elif op_kind == "delete_paragraph":
            results.append(apply_delete_paragraph_op(op, paragraph_index))
        elif op_kind == "replace_paragraph_range":
            results.append(apply_replace_paragraph_range_op(op, doc, paragraph_index))
        elif op_kind == "replace_paragraph_range_markdown":
            results.append(apply_replace_paragraph_range_markdown_op(op, doc, paragraph_index))
        else:
            raise ValueError(
                f"Unbekannte op '{op_kind}' bei Op #{i}. "
                "Unterstützt: replace_text, set_paragraph, delete_paragraph, "
                "replace_paragraph_range, replace_paragraph_range_markdown"
            )

    args.outfile.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.outfile))

    print(json.dumps({"in": str(args.infile), "out": str(args.outfile), "ops": len(results), "results": results}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(2)
