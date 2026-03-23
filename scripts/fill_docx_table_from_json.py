#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document


ALLOWED_LAYOUTS = {"cell-map"}
ALLOWED_MODES = {"replace", "append"}


def existing_docx(path_str: str) -> Path:
    p = Path(path_str)
    if not p.exists():
        raise argparse.ArgumentTypeError(f"Datei nicht gefunden: {path_str}")
    if p.suffix.lower() != ".docx":
        raise argparse.ArgumentTypeError(f"Keine .docx-Datei: {path_str}")
    return p


def existing_json(path_str: str) -> Path:
    p = Path(path_str)
    if not p.exists():
        raise argparse.ArgumentTypeError(f"JSON-Datei nicht gefunden: {path_str}")
    if p.suffix.lower() not in {".json", ".jsonl"}:
        raise argparse.ArgumentTypeError(f"Keine JSON-Datei: {path_str}")
    return p


def parse_args() -> argparse.Namespace:
    ap = argparse.ArgumentParser(description="Fill arbitrary DOCX table cells from JSON spec")
    ap.add_argument("--in", dest="infile", type=existing_docx, required=True)
    ap.add_argument("--out", dest="outfile", type=Path, required=True)
    ap.add_argument("--spec", dest="specfile", type=existing_json, required=True)
    return ap.parse_args()


def load_spec(path: Path) -> dict[str, Any]:
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Spec muss ein JSON-Objekt sein.")

    table_index = data.get("table_index")
    if not isinstance(table_index, int) or table_index < 1:
        raise ValueError("table_index muss Integer >= 1 sein.")

    layout = data.get("layout", "cell-map")
    if layout not in ALLOWED_LAYOUTS:
        raise ValueError(f"layout muss eines von {sorted(ALLOWED_LAYOUTS)} sein.")

    cells = data.get("cells")
    if not isinstance(cells, list):
        raise ValueError("Für layout='cell-map' muss 'cells' eine Liste sein.")

    normalized_cells: list[dict[str, Any]] = []
    for i, item in enumerate(cells, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"cells[{i}] muss ein Objekt sein.")

        row = item.get("row")
        col = item.get("col")
        text = item.get("text", "")
        clear_first = item.get("clear_first", False)
        mode = item.get("mode", "replace")

        if not isinstance(row, int) or row < 1:
            raise ValueError(f"cells[{i}].row muss Integer >= 1 sein.")
        if not isinstance(col, int) or col < 1:
            raise ValueError(f"cells[{i}].col muss Integer >= 1 sein.")
        if not isinstance(text, str):
            raise ValueError(f"cells[{i}].text muss String sein.")
        if not isinstance(clear_first, bool):
            raise ValueError(f"cells[{i}].clear_first muss Boolean sein.")
        if mode not in ALLOWED_MODES:
            raise ValueError(f"cells[{i}].mode muss eines von {sorted(ALLOWED_MODES)} sein.")

        normalized_cells.append(
            {
                "row": row,
                "col": col,
                "text": text,
                "clear_first": clear_first,
                "mode": mode,
            }
        )

    return {
        "table_index": table_index,
        "layout": layout,
        "cells": normalized_cells,
    }


def fill_cells(table, cells_spec: list[dict[str, Any]]) -> list[dict[str, Any]]:
    updated: list[dict[str, Any]] = []
    row_count = len(table.rows)

    for idx, cell_spec in enumerate(cells_spec, start=1):
        row_1b = cell_spec["row"]
        col_1b = cell_spec["col"]
        text = cell_spec["text"]
        clear_first = cell_spec["clear_first"]
        mode = cell_spec["mode"]

        if row_1b > row_count:
            raise ValueError(f"cells[{idx}] row={row_1b} außerhalb der Tabelle (rows={row_count}).")

        row = table.rows[row_1b - 1]
        col_count = len(row.cells)
        if col_1b > col_count:
            raise ValueError(
                f"cells[{idx}] col={col_1b} außerhalb der Zeile {row_1b} (cols={col_count})."
            )

        cell = row.cells[col_1b - 1]
        before = cell.text or ""

        if clear_first:
            cell.text = ""

        if mode == "replace":
            cell.text = text
        else:  # append
            base = cell.text or ""
            cell.text = f"{base}{text}"

        updated.append(
            {
                "cell_index": idx,
                "row": row_1b,
                "col": col_1b,
                "mode": mode,
                "clear_first": clear_first,
                "changed": before != (cell.text or ""),
            }
        )

    return updated


def main() -> int:
    args = parse_args()
    spec = load_spec(args.specfile)
    doc = Document(str(args.infile))

    table_index = spec["table_index"]
    if len(doc.tables) < table_index:
        raise ValueError(f"Dokument hat nur {len(doc.tables)} Tabellen; table_index={table_index} ist ungültig.")

    table = doc.tables[table_index - 1]
    updated = fill_cells(table, spec["cells"])

    args.outfile.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.outfile))

    print(
        json.dumps(
            {
                "ok": True,
                "in": str(args.infile),
                "out": str(args.outfile),
                "layout": spec["layout"],
                "table_index": table_index,
                "updated_cells": len(updated),
                "results": updated,
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
