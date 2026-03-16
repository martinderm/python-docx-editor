#!/usr/bin/env python3
"""Extract a .docx into structured JSON for LLM workflows.

Default output (v2): hierarchical document structure grouped by sections/headings.
Optional RAG output (v1): flat blocks + chunks via --rag-output.

Usage:
  py -3 scripts/extract_docx_for_llm.py --in input.docx --out structure.json
  py -3 scripts/extract_docx_for_llm.py --in input.docx --out structure.json --rag-output rag.json
"""

from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
import sys

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class Block:
    block_id: str
    block_type: str
    section_path: str
    text: str
    style: str | None = None
    row_range: str | None = None


def existing_docx(path_str: str) -> Path:
    p = Path(path_str)
    if not p.exists():
        raise argparse.ArgumentTypeError(f"Datei nicht gefunden: {path_str}")
    if p.suffix.lower() != ".docx":
        raise argparse.ArgumentTypeError(f"Keine .docx-Datei: {path_str}")
    return p


def iter_block_items(parent: DocumentObject):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def normalize_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    m = re.match(r"Heading\s+(\d+)", style_name)
    if m:
        return int(m.group(1))
    m = re.match(r"Überschrift\s+(\d+)", style_name)
    if m:
        return int(m.group(1))
    return None


def is_list_style(style_name: str | None) -> bool:
    if not style_name:
        return False
    s = style_name.lower()
    return any(k in s for k in ["list", "aufl", "aufz", "bullet", "number"])


def section_path_text(stack: list[str]) -> str:
    if not stack:
        return "ROOT"
    return " / ".join(stack)


def table_to_lines(table: Table) -> list[str]:
    lines: list[str] = []
    for r_idx, row in enumerate(table.rows, start=1):
        cells = [normalize_ws(c.text) for c in row.cells]
        if any(cells):
            lines.append(f"R{r_idx}: " + " | ".join(cells))
    return lines


def paragraph_numbering_info(p: Paragraph) -> dict[str, Any] | None:
    ppr = p._p.pPr  # pylint: disable=protected-access
    if ppr is None or ppr.numPr is None:
        return None
    num_id = ppr.numPr.numId
    ilvl = ppr.numPr.ilvl
    out: dict[str, Any] = {}
    if num_id is not None and num_id.val is not None:
        out["num_id"] = int(num_id.val)
    if ilvl is not None and ilvl.val is not None:
        out["level"] = int(ilvl.val)
    return out or None


def extract_v2(source: Path, doc: DocumentObject, table_row_batch: int) -> dict:
    p_i = 0
    t_i = 0
    h_i = 0
    f_i = 0

    section_stack: list[dict[str, Any]] = []
    top_sections: list[dict[str, Any]] = []
    pre_heading_content: list[dict[str, Any]] = []

    def push_section(level: int, title: str, block_id: str, style: str | None) -> dict[str, Any]:
        while section_stack and section_stack[-1]["level"] >= level:
            section_stack.pop()

        node = {
            "id": f"sec_{len(top_sections) + 1 + sum(len(s.get('children', [])) for s in top_sections)}",
            "level": level,
            "title": title,
            "block_id": block_id,
            "style": style,
            "content": [],
            "children": [],
        }

        if section_stack:
            section_stack[-1]["children"].append(node)
        else:
            top_sections.append(node)
        section_stack.append(node)
        return node

    def current_content_target() -> list[dict[str, Any]]:
        if section_stack:
            return section_stack[-1]["content"]
        return pre_heading_content

    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = normalize_ws(item.text)
            if not text:
                continue

            p_i += 1
            bid = f"p_{p_i}"
            style_name = item.style.name if item.style else None
            lvl = heading_level(style_name)

            if lvl is not None:
                push_section(level=lvl, title=text, block_id=bid, style=style_name)
                continue

            node_type = "list_item" if (is_list_style(style_name) or paragraph_numbering_info(item)) else "paragraph"
            node: dict[str, Any] = {
                "type": node_type,
                "block_id": bid,
                "text": text,
                "style": style_name,
            }
            num_info = paragraph_numbering_info(item)
            if num_info:
                node["list"] = num_info
            current_content_target().append(node)

        elif isinstance(item, Table):
            t_i += 1
            lines = table_to_lines(item)
            if not lines:
                continue
            if table_row_batch <= 0:
                table_row_batch = len(lines)

            table_node = {
                "type": "table",
                "table_index": t_i,
                "row_count": len(item.rows),
                "col_count": len(item.columns),
                "parts": [],
            }

            for start in range(0, len(lines), table_row_batch):
                end = min(start + table_row_batch, len(lines))
                row_range = f"{start + 1}-{end}"
                bid = f"t_{t_i}_r{start + 1}_{end}"
                table_node["parts"].append(
                    {
                        "block_id": bid,
                        "row_range": row_range,
                        "text": "\n".join(lines[start:end]),
                    }
                )

            current_content_target().append(table_node)

    headers: list[dict[str, Any]] = []
    footers: list[dict[str, Any]] = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            text = normalize_ws(p.text)
            if not text:
                continue
            h_i += 1
            headers.append(
                {
                    "type": "header",
                    "block_id": f"h_{h_i}",
                    "text": text,
                    "style": p.style.name if p.style else None,
                }
            )
        for p in section.footer.paragraphs:
            text = normalize_ws(p.text)
            if not text:
                continue
            f_i += 1
            footers.append(
                {
                    "type": "footer",
                    "block_id": f"f_{f_i}",
                    "text": text,
                    "style": p.style.name if p.style else None,
                }
            )

    def count_nodes_content(content: list[dict[str, Any]]) -> dict[str, int]:
        counts = {"paragraph": 0, "list_item": 0, "table": 0}
        for n in content:
            t = n.get("type")
            if t in counts:
                counts[t] += 1
        return counts

    def walk_sections(sections: list[dict[str, Any]]) -> tuple[int, dict[str, int]]:
        sec_count = 0
        totals = {"paragraph": 0, "list_item": 0, "table": 0}
        for s in sections:
            sec_count += 1
            c = count_nodes_content(s["content"])
            for k, v in c.items():
                totals[k] += v
            child_count, child_totals = walk_sections(s["children"])
            sec_count += child_count
            for k, v in child_totals.items():
                totals[k] += v
        return sec_count, totals

    sec_total, typed_totals = walk_sections(top_sections)
    pre_counts = count_nodes_content(pre_heading_content)
    for k, v in pre_counts.items():
        typed_totals[k] += v

    return {
        "schema": "docx-structure.v2",
        "source": str(source),
        "stats": {
            "sections": sec_total,
            "content": typed_totals,
            "headers": len(headers),
            "footers": len(footers),
            "pre_heading_nodes": len(pre_heading_content),
        },
        "document": {
            "pre_heading": pre_heading_content,
            "sections": top_sections,
            "headers": headers,
            "footers": footers,
        },
    }


# ------- optional v1 RAG output -------

def extract_v1_blocks(doc: DocumentObject, table_row_batch: int) -> list[Block]:
    blocks: list[Block] = []
    heading_stack: list[str] = []

    para_i = 0
    table_i = 0
    header_i = 0
    footer_i = 0

    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = normalize_ws(item.text)
            if not text:
                continue
            para_i += 1

            style_name = item.style.name if item.style else None
            lvl = heading_level(style_name)
            if lvl is not None:
                while len(heading_stack) >= lvl:
                    heading_stack.pop()
                heading_stack.append(text)
                blocks.append(
                    Block(
                        block_id=f"p_{para_i}",
                        block_type="heading",
                        section_path=section_path_text(heading_stack),
                        text=text,
                        style=style_name,
                    )
                )
                continue

            blocks.append(
                Block(
                    block_id=f"p_{para_i}",
                    block_type="paragraph",
                    section_path=section_path_text(heading_stack),
                    text=text,
                    style=style_name,
                )
            )

        elif isinstance(item, Table):
            table_i += 1
            lines = table_to_lines(item)
            if not lines:
                continue
            if table_row_batch <= 0:
                table_row_batch = len(lines)
            for start in range(0, len(lines), table_row_batch):
                end = min(start + table_row_batch, len(lines))
                row_range = f"{start + 1}-{end}"
                table_text = "\n".join(lines[start:end])
                blocks.append(
                    Block(
                        block_id=f"t_{table_i}_r{start + 1}_{end}",
                        block_type="table",
                        section_path=section_path_text(heading_stack),
                        text=table_text,
                        row_range=row_range,
                    )
                )

    for section in doc.sections:
        for p in section.header.paragraphs:
            text = normalize_ws(p.text)
            if not text:
                continue
            header_i += 1
            blocks.append(
                Block(
                    block_id=f"h_{header_i}",
                    block_type="header",
                    section_path="HEADER",
                    text=text,
                    style=p.style.name if p.style else None,
                )
            )
        for p in section.footer.paragraphs:
            text = normalize_ws(p.text)
            if not text:
                continue
            footer_i += 1
            blocks.append(
                Block(
                    block_id=f"f_{footer_i}",
                    block_type="footer",
                    section_path="FOOTER",
                    text=text,
                    style=p.style.name if p.style else None,
                )
            )

    return blocks


def chunk_blocks(blocks: list[Block], max_chars: int, overlap: int) -> list[dict]:
    chunks: list[dict] = []
    idx = 0
    while idx < len(blocks):
        current: list[Block] = []
        size = 0
        start_idx = idx

        while idx < len(blocks):
            b = blocks[idx]
            b_len = len(b.text) + 32
            if current and size + b_len > max_chars:
                break
            current.append(b)
            size += b_len
            idx += 1

        if not current:
            current = [blocks[idx]]
            idx += 1

        text = "\n\n".join(f"[{b.block_id}] {b.text}" for b in current)
        chunks.append(
            {
                "chunk_id": f"c_{len(chunks) + 1:04d}",
                "block_ids": [b.block_id for b in current],
                "section_paths": sorted({b.section_path for b in current}),
                "block_types": sorted({b.block_type for b in current}),
                "text": text,
                "char_count": len(text),
            }
        )

        if overlap > 0 and idx < len(blocks):
            idx = max(start_idx + 1, idx - overlap)

    return chunks


def build_v1(source: Path, blocks: list[Block], chunks: list[dict]) -> dict:
    return {
        "schema": "docx-llm-chunks.v1",
        "source": str(source),
        "stats": {
            "blocks": len(blocks),
            "chunks": len(chunks),
            "by_type": {
                "heading": sum(1 for b in blocks if b.block_type == "heading"),
                "paragraph": sum(1 for b in blocks if b.block_type == "paragraph"),
                "table": sum(1 for b in blocks if b.block_type == "table"),
                "header": sum(1 for b in blocks if b.block_type == "header"),
                "footer": sum(1 for b in blocks if b.block_type == "footer"),
            },
        },
        "blocks": [
            {
                "block_id": b.block_id,
                "type": b.block_type,
                "section_path": b.section_path,
                "style": b.style,
                "row_range": b.row_range,
                "text": b.text,
            }
            for b in blocks
        ],
        "chunks": chunks,
    }


def write_json(path: Path, data: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Extract DOCX as structured JSON")
    p.add_argument("--in", dest="infile", type=existing_docx, required=True)
    p.add_argument("--out", dest="outfile", type=Path, required=True, help="Primary output: v2 structure")
    p.add_argument("--table-row-batch", type=int, default=25, help="Rows per table block/part")

    # optional v1/RAG compatibility output
    p.add_argument("--rag-output", dest="rag_outfile", type=Path, help="Optional v1 (blocks+chunks) output path")
    p.add_argument("--max-chars", type=int, default=4000, help="Target chunk size for --rag-output")
    p.add_argument("--overlap-blocks", type=int, default=1, help="Block overlap for --rag-output")
    return p.parse_args()


def main() -> int:
    args = parse_args()
    doc = Document(str(args.infile))

    v2 = extract_v2(args.infile, doc, table_row_batch=args.table_row_batch)
    write_json(args.outfile, v2)

    result: dict[str, Any] = {
        "out": str(args.outfile),
        "schema": v2["schema"],
        "sections": v2["stats"]["sections"],
    }

    if args.rag_outfile:
        blocks = extract_v1_blocks(doc, table_row_batch=args.table_row_batch)
        chunks = chunk_blocks(blocks, max_chars=max(500, args.max_chars), overlap=max(0, args.overlap_blocks))
        v1 = build_v1(args.infile, blocks, chunks)
        write_json(args.rag_outfile, v1)
        result["rag_output"] = str(args.rag_outfile)
        result["rag_blocks"] = v1["stats"]["blocks"]
        result["rag_chunks"] = v1["stats"]["chunks"]

    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(2)
