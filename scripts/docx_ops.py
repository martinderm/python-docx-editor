#!/usr/bin/env python3
"""Small CLI helpers for .docx operations with python-docx.

Usage examples:
  py -3 scripts/docx_ops.py text --in input.docx
  py -3 scripts/docx_ops.py stats --in input.docx
  py -3 scripts/docx_ops.py replace --in in.docx --out out.docx --find "Alt" --replace "Neu"
  py -3 scripts/docx_ops.py inspect --in input.docx [--json]
"""

from __future__ import annotations

import argparse
import json
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document


def cmd_text(infile: Path) -> None:
    doc = Document(str(infile))
    lines: list[str] = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    print("\n".join(lines))


def cmd_stats(infile: Path) -> None:
    doc = Document(str(infile))

    paragraph_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    image_count = 0

    for rel in doc.part._rels.values():  # pylint: disable=protected-access
        if "image" in rel.reltype:
            image_count += 1

    data = {
        "file": str(infile),
        "paragraphs": paragraph_count,
        "tables": table_count,
        "images": image_count,
    }
    print(json.dumps(data, ensure_ascii=False, indent=2))


def cmd_replace(infile: Path, outfile: Path, old: str, new: str) -> None:
    doc = Document(str(infile))
    replaced = 0

    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
            replaced += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old in cell.text:
                    cell.text = cell.text.replace(old, new)
                    replaced += 1

    doc.save(str(outfile))
    print(json.dumps({"out": str(outfile), "replaced_blocks": replaced}, ensure_ascii=False))


def inspect_formatting(infile: Path) -> dict:
    namespaces = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    results = {
        "comments": [],
        "revisions": {
            "insertions": [],
            "deletions": []
        },
        "highlights": []
    }

    if not infile.exists():
        return results

    with zipfile.ZipFile(str(infile)) as z:
        namelist = z.namelist()

        # 1. Parse comments
        if 'word/comments.xml' in namelist:
            comments_xml = z.read('word/comments.xml')
            root = ET.fromstring(comments_xml)
            for comment in root.findall('.//w:comment', namespaces):
                author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
                date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date')
                text_elems = comment.findall('.//w:t', namespaces)
                text = "".join(t.text for t in text_elems if t.text)
                results["comments"].append({
                    "author": author or "Unknown",
                    "date": date or "Unknown",
                    "text": text
                })

        # 2. Parse main document.xml for revisions and highlights/shading
        if 'word/document.xml' in namelist:
            doc_xml = z.read('word/document.xml')
            root = ET.fromstring(doc_xml)

            # Revisions: Insertions
            for ins in root.findall('.//w:ins', namespaces):
                author = ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
                date = ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date')
                text_elems = ins.findall('.//w:t', namespaces)
                text = "".join(t.text for t in text_elems if t.text)
                if text.strip():
                    results["revisions"]["insertions"].append({
                        "author": author or "Unknown",
                        "date": date or "Unknown",
                        "text": text
                    })

            # Revisions: Deletions
            for del_elem in root.findall('.//w:delText', namespaces):
                text = del_elem.text or ""
                if text.strip():
                    results["revisions"]["deletions"].append({
                        "text": text
                    })

            # Traverse the body elements in document order to track nearest heading and positions
            body = root.find('.//w:body', namespaces)
            if body is not None:
                p_index = 0
                tbl_index = 0
                nearest_heading = "None"

                def inspect_paragraph_element(p, pos_info):
                    nonlocal nearest_heading
                    pPr = p.find('.//w:pPr', namespaces)
                    p_has_format = False
                    p_fmt_details = {}

                    style_val = ""
                    if pPr is not None:
                        style_elem = pPr.find('.//w:pStyle', namespaces)
                        if style_elem is not None:
                            style_val = style_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') or ""

                        # Paragraph-level shading
                        p_shd = pPr.find('.//w:shd', namespaces)
                        if p_shd is not None:
                            fill = p_shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                            if fill and fill.lower() not in ['auto', 'clear', 'ffffff']:
                                p_has_format = True
                                p_fmt_details["paragraph_shading"] = fill

                        # Paragraph-level default run properties
                        pPr_rPr = pPr.find('.//w:rPr', namespaces)
                        if pPr_rPr is not None:
                            p_color = pPr_rPr.find('.//w:color', namespaces)
                            p_hl = pPr_rPr.find('.//w:highlight', namespaces)
                            if p_color is not None:
                                val = p_color.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if val and val.lower() not in ['auto', '000000', 'windowtext']:
                                    p_has_format = True
                                    p_fmt_details["paragraph_color"] = val
                            if p_hl is not None:
                                val = p_hl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if val and val.lower() not in ['none']:
                                    p_has_format = True
                                    p_fmt_details["paragraph_highlight"] = val

                    text_elems = p.findall('.//w:t', namespaces)
                    text = "".join(t.text for t in text_elems if t.text)

                    # Update nearest heading if this is a heading style (or title/subtitle)
                    if style_val and text.strip():
                        if any(q in style_val.lower() for q in ["heading", "titel", "title", "heading"]):
                            nearest_heading = text.strip()

                    # Record paragraph-level findings
                    if p_has_format and text.strip():
                        pos = pos_info.copy()
                        pos["nearest_heading"] = nearest_heading
                        results["highlights"].append({
                            "text": text.strip(),
                            "position": pos,
                            "details": p_fmt_details
                        })

                    # Record run-level findings
                    for r_idx, r in enumerate(p.findall('.//w:r', namespaces)):
                        rPr = r.find('.//w:rPr', namespaces)
                        if rPr is not None:
                            color = rPr.find('.//w:color', namespaces)
                            highlight = rPr.find('.//w:highlight', namespaces)
                            shading = rPr.find('.//w:shd', namespaces)

                            t_elem = r.find('.//w:t', namespaces)
                            r_text = t_elem.text if t_elem is not None else ""

                            if not r_text.strip():
                                continue

                            r_has_format = False
                            r_fmt_details = {}

                            if color is not None:
                                val = color.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if val and val.lower() not in ['auto', '000000', 'windowtext']:
                                    r_has_format = True
                                    r_fmt_details["color"] = val
                            if highlight is not None:
                                val = highlight.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if val and val.lower() not in ['none']:
                                    r_has_format = True
                                    r_fmt_details["highlight"] = val
                            if shading is not None:
                                fill = shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                                if fill and fill.lower() not in ['auto', 'clear', 'ffffff', 'none']:
                                    r_has_format = True
                                    r_fmt_details["shading"] = fill

                            if r_has_format:
                                pos = pos_info.copy()
                                pos["nearest_heading"] = nearest_heading
                                pos["run_index"] = r_idx
                                results["highlights"].append({
                                    "text": r_text.strip(),
                                    "position": pos,
                                    "details": r_fmt_details
                                })

                # Traverse children in document order
                for child in body:
                    tag = child.tag.split('}')[-1]
                    if tag == 'p':
                        pos_info = {
                            "location_type": "Main Text",
                            "paragraph_index": p_index
                        }
                        inspect_paragraph_element(child, pos_info)
                        p_index += 1
                    elif tag == 'tbl':
                        # Resolve table heading by looking at the first cell of the first row
                        first_row = child.find('.//w:tr', namespaces)
                        table_heading = f"Table {tbl_index}"
                        if first_row is not None:
                            first_cell = first_row.find('.//w:tc', namespaces)
                            if first_cell is not None:
                                tc_text = "".join(t.text for t in first_cell.findall('.//w:t', namespaces) if t.text)
                                if tc_text.strip():
                                    table_heading = tc_text.strip()

                        old_heading = nearest_heading
                        nearest_heading = table_heading

                        for r_idx, row in enumerate(child.findall('.//w:tr', namespaces)):
                            for c_idx, cell in enumerate(row.findall('.//w:tc', namespaces)):
                                for cell_p_idx, cell_p in enumerate(cell.findall('.//w:p', namespaces)):
                                    pos_info = {
                                        "location_type": "Table Cell",
                                        "table_index": tbl_index,
                                        "row_index": r_idx,
                                        "col_index": c_idx,
                                        "cell_paragraph_index": cell_p_idx
                                    }
                                    inspect_paragraph_element(cell_p, pos_info)

                        nearest_heading = old_heading
                        tbl_index += 1

    return results


def cmd_inspect(infile: Path, output_json: bool = False) -> None:
    results = inspect_formatting(infile)
    if output_json:
        print(json.dumps(results, ensure_ascii=False, indent=2))
        return

    print(f"=== Formatting Inspection for {infile.name} ===")

    comments = results["comments"]
    if comments:
        print(f"\nComments ({len(comments)} found):")
        for c in comments:
            print(f"  [{c['author']} @ {c['date']}]: '{c['text']}'")
    else:
        print("\nNo comments found.")

    ins = results["revisions"]["insertions"]
    dels = results["revisions"]["deletions"]
    if ins or dels:
        print("\nTracked Revisions:")
        if ins:
            print(f"  Insertions ({len(ins)}):")
            for item in ins:
                print(f"    [{item['author']} @ {item['date']}]: '{item['text']}'")
        if dels:
            print(f"  Deletions ({len(dels)}):")
            for item in dels:
                print(f"    '{item['text']}'")
    else:
        print("\nNo tracked revisions found.")

    highlights = results["highlights"]
    if highlights:
        print(f"\nHighlighted / Colored Text ({len(highlights)} instances):")
        for hl in highlights:
            fmt_desc = ", ".join(f"{k}: {v}" for k, v in hl["details"].items())
            pos = hl["position"]
            if pos["location_type"] == "Table Cell":
                loc_desc = f"Table {pos['table_index']}, R{pos['row_index']}C{pos['col_index']}P{pos['cell_paragraph_index']}"
            else:
                loc_desc = f"Para {pos['paragraph_index']}"
            run_desc = f", Run {pos['run_index']}" if "run_index" in pos else ""
            
            print(f"  [{pos['nearest_heading']} -> {loc_desc}{run_desc}] '{hl['text']}' ({fmt_desc})")
    else:
        print("\nNo highlighted / colored text found.")


def existing_docx(path_str: str) -> Path:
    p = Path(path_str)
    if not p.exists():
        raise argparse.ArgumentTypeError(f"Datei nicht gefunden: {path_str}")
    if p.suffix.lower() != ".docx":
        raise argparse.ArgumentTypeError(f"Keine .docx-Datei: {path_str}")
    return p


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="DOCX Ops via python-docx")
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_text = sub.add_parser("text", help="Text aus DOCX extrahieren")
    p_text.add_argument("--in", dest="infile", type=existing_docx, required=True)

    p_stats = sub.add_parser("stats", help="Basis-Statistik als JSON")
    p_stats.add_argument("--in", dest="infile", type=existing_docx, required=True)

    p_replace = sub.add_parser("replace", help="Einfache Text-Ersetzung")
    p_replace.add_argument("--in", dest="infile", type=existing_docx, required=True)
    p_replace.add_argument("--out", dest="outfile", type=Path, required=True)
    p_replace.add_argument("--find", dest="old", required=True)
    p_replace.add_argument("--replace", dest="new", required=True)

    p_inspect = sub.add_parser("inspect", help="Word-Kommentare, Revisionen und farbige Markierungen auslesen")
    p_inspect.add_argument("--in", dest="infile", type=existing_docx, required=True)
    p_inspect.add_argument("--json", action="store_true", help="Ergebnisse als JSON ausgeben")

    return parser.parse_args()


def main() -> int:
    args = parse_args()
    if args.cmd == "text":
        cmd_text(args.infile)
    elif args.cmd == "stats":
        cmd_stats(args.infile)
    elif args.cmd == "replace":
        cmd_replace(args.infile, args.outfile, args.old, args.new)
    elif args.cmd == "inspect":
        cmd_inspect(args.infile, args.json)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(2)
